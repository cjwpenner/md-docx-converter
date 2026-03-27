from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
import markdown_it
import shutil
import tempfile
import zipfile
from md_docx_converter.heading_mapper import md_heading_offset
from md_docx_converter.image_handler import resolve_image_path, embed_image

_DOTM_CT = "application/vnd.ms-word.template.macroEnabledTemplate.main+xml"
_DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"


def _open_template(template_path: Path) -> Document:
    """
    Open a Word template as a python-docx Document with all standard styles available.

    Normal.dotm only defines the 'Normal' style in its XML — all other built-in
    styles (Title, Heading 1, etc.) are inherited from Word's built-in style dictionary
    but are absent from the file. python-docx can't apply styles it hasn't seen.

    Strategy: open python-docx's built-in default document (which has all standard
    styles), then copy the Normal paragraph formatting from the .dotm so the user's
    font/spacing preferences are applied to the Normal style.
    """
    doc = Document()  # opens built-in default template — has all standard styles

    # Try to extract Normal style properties from the .dotm and apply them
    suffix = template_path.suffix.lower()
    if suffix in (".dotm", ".dotx", ".docx"):
        try:
            tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            tmp.close()
            tmp_path = Path(tmp.name)
            shutil.copy2(template_path, tmp_path)
            with zipfile.ZipFile(str(tmp_path), "r") as zin:
                files = {n: zin.read(n) for n in zin.namelist()}
            ct_xml = files["[Content_Types].xml"].decode("utf-8")
            ct_xml = ct_xml.replace(_DOTM_CT, _DOCX_CT)
            files["[Content_Types].xml"] = ct_xml.encode("utf-8")
            with zipfile.ZipFile(str(tmp_path), "w", zipfile.ZIP_DEFLATED) as zout:
                for name, data in files.items():
                    zout.writestr(name, data)
            tmpl_doc = Document(str(tmp_path))
            tmp_path.unlink()
            # Copy Normal style's font from template to our document
            tmpl_normal = tmpl_doc.styles["Normal"]
            our_normal = doc.styles["Normal"]
            if tmpl_normal.font.name:
                our_normal.font.name = tmpl_normal.font.name
            if tmpl_normal.font.size:
                our_normal.font.size = tmpl_normal.font.size
        except Exception:
            pass  # If anything fails, just use the built-in default as-is

    return doc


def convert_md_to_docx(md_path: Path, out_path: Path, template_path: Path) -> Path:
    md_text = md_path.read_text(encoding="utf-8", errors="replace")
    offset = md_heading_offset(md_text)
    md_dir = md_path.parent

    doc = _open_template(template_path)
    # Remove all existing body content from template (keep sectPr for page layout)
    for element in list(doc.element.body):
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag in ('p', 'tbl'):
            doc.element.body.remove(element)

    mdi = (
        markdown_it.MarkdownIt("commonmark")
        .enable("table")
        .enable("strikethrough")
    )
    tokens = mdi.parse(md_text)

    _write_tokens(doc, tokens, offset, md_dir)

    doc.save(str(out_path))
    return out_path


def _set_style(para, style_name: str):
    """Apply a named style to a paragraph, falling back to Normal if not found."""
    try:
        para.style = style_name
    except KeyError:
        pass  # Style not in template — leave as Normal


def _add_para(doc, style_name: str):
    """Add a paragraph and apply style, with fallback."""
    para = doc.add_paragraph()
    _set_style(para, style_name)
    return para


def _write_tokens(doc, tokens, offset, md_dir):
    i = 0
    while i < len(tokens):
        tok = tokens[i]

        if tok.type == "heading_open":
            level = int(tok.tag[1])  # h1→1, h2→2, etc.
            inline = tokens[i + 1]
            style = _heading_style(level, offset)
            para = _add_para(doc, style)
            _apply_inline(para, inline.children or [], md_dir)
            i += 3  # heading_open, inline, heading_close

        elif tok.type == "paragraph_open":
            inline = tokens[i + 1]
            para = _add_para(doc, "Normal")
            _apply_inline(para, inline.children or [], md_dir)
            i += 3

        elif tok.type == "bullet_list_open":
            i, _ = _write_list(doc, tokens, i, "List Bullet", depth=0, md_dir=md_dir)

        elif tok.type == "ordered_list_open":
            i, _ = _write_list(doc, tokens, i, "List Number", depth=0, md_dir=md_dir)

        elif tok.type == "blockquote_open":
            i = _write_blockquote(doc, tokens, i, md_dir)

        elif tok.type == "fence":
            _write_code_block(doc, tok)
            i += 1

        elif tok.type == "hr":
            _write_hr(doc)
            i += 1

        elif tok.type in ("html_block", "html_inline"):
            i += 1  # silently drop HTML

        elif tok.type == "table_open":
            i = _write_table(doc, tokens, i, md_dir)

        else:
            i += 1


def _heading_style(level: int, offset: int) -> str:
    if offset == 1 and level == 1:
        return "Title"
    adjusted = level - offset
    adjusted = max(1, min(adjusted, 9))
    return f"Heading {adjusted}"


def _apply_inline(para, children, md_dir: Path):
    """Apply inline tokens to a paragraph as runs, handling bold/italic/code/images/links."""
    bold = False
    italic = False
    strike = False
    link_href = None

    i = 0
    while i < len(children):
        child = children[i]


        if child.type == "strong_open":
            bold = True
        elif child.type == "strong_close":
            bold = False
        elif child.type == "em_open":
            italic = True
        elif child.type == "em_close":
            italic = False
        elif child.type == "s_open":
            strike = True
        elif child.type == "s_close":
            strike = False
        elif child.type == "link_open":
            link_href = child.attrs.get("href", "")
        elif child.type == "link_close":
            link_href = None
        elif child.type == "code_inline":
            run = para.add_run(f"'{child.content}'")
        elif child.type == "image":
            src = child.attrs.get("src", "")
            resolved = resolve_image_path(src, md_dir)
            if resolved:
                embed_image(para, resolved)
            else:
                para.add_run(f"[image not found: {src}]")
        elif child.type in ("softbreak", "hardbreak"):
            para.add_run("\n")
        elif child.type == "text":
            content = child.content
            # If inside a link, append the URL as plain text after link text
            if link_href and i + 1 < len(children) and children[i + 1].type == "link_close":
                content = f"{content} ({link_href})"
            run = para.add_run(content)
            run.bold = bold
            run.italic = italic
            run.font.strike = strike
        elif child.type == "html_inline":
            pass  # drop inline HTML

        i += 1


def _task_list_prefix(children) -> tuple[str, list]:
    """
    If inline children start with text '[ ] ' or '[x] ', return the checkbox
    character and the children list with the prefix stripped from the first text token.
    Returns ('', children) if not a task list item.
    """
    if not children:
        return "", children
    first = children[0]
    if first.type != "text":
        return "", children

    import markdown_it.token as mit
    for prefix, symbol in [("[ ] ", "☐ "), ("[x] ", "☑ "), ("[X] ", "☑ ")]:
        if first.content.startswith(prefix):
            new_first = mit.Token(first.type, first.tag, first.nesting)
            new_first.content = first.content[len(prefix):]
            return symbol, [new_first] + list(children[1:])

    return "", children


def _write_list(doc, tokens, i, base_style, depth, md_dir):
    """Write a list block. Returns (new_index, None)."""
    open_type = tokens[i].type
    close_type = open_type.replace("_open", "_close")
    i += 1
    while i < len(tokens):
        tok = tokens[i]
        if tok.type == close_type:
            i += 1
            break
        elif tok.type == "list_item_open":
            i += 1
            while i < len(tokens) and tokens[i].type != "list_item_close":
                inner = tokens[i]
                if inner.type == "paragraph_open":
                    inline = tokens[i + 1]
                    children = inline.children or []
                    style = base_style if depth == 0 else f"{base_style} {min(depth + 1, 3)}"
                    checkbox, children = _task_list_prefix(children)
                    para = _add_para(doc, style)
                    if checkbox:
                        para.add_run(checkbox)
                    _apply_inline(para, children, md_dir)
                    i += 3  # paragraph_open, inline, paragraph_close
                elif inner.type in ("bullet_list_open", "ordered_list_open"):
                    nested_style = "List Bullet" if "bullet" in inner.type else "List Number"
                    i, _ = _write_list(doc, tokens, i, nested_style, depth + 1, md_dir)
                else:
                    i += 1
        else:
            i += 1
    return i, None


def _write_blockquote(doc, tokens, i, md_dir):
    i += 1  # skip blockquote_open
    depth = 1
    while i < len(tokens):
        tok = tokens[i]
        if tok.type == "blockquote_open":
            depth += 1
            i += 1
        elif tok.type == "blockquote_close":
            depth -= 1
            i += 1
            if depth == 0:
                break
        elif tok.type == "paragraph_open":
            inline = tokens[i + 1]
            para = _add_para(doc, "Quote")
            _apply_inline(para, inline.children or [], md_dir)
            i += 3
        else:
            i += 1
    return i


def _write_code_block(doc, tok):
    content = tok.content.rstrip("\n")
    para = _add_para(doc, "Normal")
    run = para.add_run(content)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    # Try to apply a Code style if it exists in the template
    try:
        para.style = doc.styles["Code"]
    except KeyError:
        pass  # No Code style in template — monospaced Normal is acceptable


def _write_hr(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _write_table(doc, tokens, i, md_dir):
    rows = []  # list of rows; each row is a list of inline-children lists
    i += 1  # skip table_open
    while i < len(tokens) and tokens[i].type != "table_close":
        tok = tokens[i]
        if tok.type in ("thead_open", "tbody_open", "thead_close", "tbody_close"):
            i += 1
        elif tok.type == "tr_open":
            row_cells = []
            i += 1
            while i < len(tokens) and tokens[i].type != "tr_close":
                if tokens[i].type in ("th_open", "td_open"):
                    inline = tokens[i + 1]
                    row_cells.append(inline.children or [])
                    i += 3  # th/td open, inline, th/td close
                else:
                    i += 1
            rows.append(row_cells)
            i += 1  # skip tr_close
        else:
            i += 1
    i += 1  # skip table_close

    if not rows:
        return i

    num_cols = max(len(r) for r in rows)
    rows = [r + [[]] * (num_cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=num_cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for r_idx, row in enumerate(rows):
        for c_idx, children in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            para = cell.paragraphs[0]
            _apply_inline(para, children, md_dir)

    return i
