from pathlib import Path
from docx import Document
from md_docx_converter.heading_mapper import docx_heading_offset
from md_docx_converter.image_handler import extract_docx_images

# Word namespace URIs
_W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_WP  = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_A   = "http://schemas.openxmlformats.org/drawingml/2006/main"
_PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
_R   = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

# Monospaced font names (lowercase)
_MONO_FONTS = {"courier new", "courier", "consolas", "lucida console",
               "monaco", "menlo", "source code pro"}


def convert_docx_to_md(docx_path: Path, out_md_path: Path) -> Path:
    offset = docx_heading_offset(docx_path)
    image_map = extract_docx_images(docx_path, out_md_path)

    doc = Document(str(docx_path))

    # Build relationship ID → extracted image Path
    rel_to_path: dict[str, Path] = {}
    for rId, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            target = rel.target_ref          # e.g. '../media/image1.png'
            internal = "word/" + target.lstrip("../")
            if internal in image_map:
                rel_to_path[rId] = image_map[internal]

    lines: list[str] = []

    for child in doc.element.body:
        tag = child.tag.split("}")[1] if "}" in child.tag else child.tag
        if tag == "p":
            para = _find_para(doc, child)
            if para is not None:
                md_line = _para_to_md(para, offset, out_md_path, rel_to_path, doc)
                if md_line is not None:
                    lines.append(md_line)
        elif tag == "tbl":
            table = _find_table(doc, child)
            if table is not None:
                lines.append(_table_to_md(table))

    out_md_path.write_text("\n\n".join(lines) + "\n", encoding="utf-8")
    return out_md_path


def _find_para(doc, elem):
    for para in doc.paragraphs:
        if para._element is elem:
            return para
    return None


def _find_table(doc, elem):
    for table in doc.tables:
        if table._element is elem:
            return table
    return None


def _para_to_md(para, offset: int, out_md_path: Path,
                rel_to_path: dict, doc) -> str | None:
    style = para.style.name

    # Code style → fenced code block
    if style == "Code":
        return f"```\n{para.text}\n```"

    text = _runs_to_md(para, out_md_path, rel_to_path)

    # Check for hyperlinks in paragraph XML
    text = _extract_hyperlinks(para, text, doc)

    if not text.strip():
        return None

    if style == "Title":
        return f"# {text}"

    if style.startswith("Heading "):
        try:
            level = int(style.split()[-1])
        except ValueError:
            return text
        md_level = level + offset
        return "#" * md_level + f" {text}"

    if style.startswith("List Bullet"):
        depth = _list_depth(style)
        return "  " * depth + f"- {text}"

    if style.startswith("List Number"):
        depth = _list_depth(style)
        return "  " * depth + f"1. {text}"

    if style in ("Quote", "Block Text", "Quotations", "Intense Quote"):
        return f"> {text}"

    return text


def _list_depth(style_name: str) -> int:
    parts = style_name.split()
    if len(parts) >= 3:
        try:
            return int(parts[-1]) - 1
        except ValueError:
            pass
    return 0


def _runs_to_md(para, out_md_path: Path, rel_to_path: dict) -> str:
    parts = []
    for run in para.runs:
        text = run.text

        # Check for inline image drawing in this run
        img_md = _run_image_md(run, out_md_path, rel_to_path)
        if img_md:
            parts.append(img_md)
            continue

        if not text:
            continue

        is_bold = bool(run.bold)
        is_italic = bool(run.italic)
        is_strike = bool(run.font.strike)
        is_mono = (run.font.name or "").lower() in _MONO_FONTS

        # Approximate unsupported formatting as bold
        is_underline = bool(run.underline)
        is_highlight = bool(run.font.highlight_color)
        is_small_caps = bool(run.font.small_caps)
        is_large = bool(run.font.size and run.font.size.pt > 14) if run.font.size else False
        approx_bold = is_underline or is_highlight or is_small_caps or is_large
        effective_bold = is_bold or approx_bold

        if is_mono:
            parts.append(f"`{text}`")
            continue

        if is_strike:
            text = f"~~{text}~~"
        if effective_bold and is_italic:
            text = f"***{text}***"
        elif effective_bold:
            text = f"**{text}**"
        elif is_italic:
            text = f"*{text}*"

        parts.append(text)

    return "".join(parts)


def _run_image_md(run, out_md_path: Path, rel_to_path: dict) -> str | None:
    """
    Check if a run contains an inline Word drawing (image).
    Word inline images: w:r/w:drawing/wp:inline/a:graphic/a:graphicData/
                        pic:pic/pic:blipFill/a:blip[@r:embed]
    """
    drawing_elems = run._element.findall(f"{{{_W}}}drawing")
    for drawing in drawing_elems:
        blip = drawing.find(
            f"{{{_WP}}}inline/{{{_A}}}graphic/{{{_A}}}graphicData"
            f"/{{{_PIC}}}pic/{{{_PIC}}}blipFill/{{{_A}}}blip"
        )
        if blip is None:
            continue
        r_embed = blip.get(f"{{{_R}}}embed")
        if r_embed and r_embed in rel_to_path:
            img_path = rel_to_path[r_embed]
            try:
                rel = img_path.relative_to(out_md_path.parent)
                return f"![]({rel.as_posix()})"
            except ValueError:
                return f"![]({img_path.as_posix()})"
    return None


def _extract_hyperlinks(para, plain_text: str, doc) -> str:
    """
    Walk paragraph XML for <w:hyperlink> elements, extract URL from relationships,
    and emit [text](url) GFM syntax. Returns plain_text unchanged if no hyperlinks.
    """
    hyperlink_tag = f"{{{_W}}}hyperlink"
    hyperlinks = para._element.findall(hyperlink_tag)
    if not hyperlinks:
        return plain_text

    parts = []
    for child in para._element:
        if child.tag == hyperlink_tag:
            run_texts = []
            for r_elem in child.findall(f"{{{_W}}}r"):
                t_elem = r_elem.find(f"{{{_W}}}t")
                if t_elem is not None and t_elem.text:
                    run_texts.append(t_elem.text)
            link_text = "".join(run_texts)

            r_id = child.get(f"{{{_R}}}id")
            url = ""
            if r_id and r_id in doc.part.rels:
                url = doc.part.rels[r_id].target_ref

            parts.append(f"[{link_text}]({url})" if url else link_text)
        elif child.tag == f"{{{_W}}}r":
            t_elem = child.find(f"{{{_W}}}t")
            if t_elem is not None and t_elem.text:
                parts.append(t_elem.text)

    rebuilt = "".join(parts)
    return rebuilt if rebuilt.strip() else plain_text


def _table_to_md(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")

    if not rows:
        return ""

    num_cols = len(table.columns)
    separator = "| " + " | ".join(["---"] * num_cols) + " |"
    rows.insert(1, separator)
    return "\n".join(rows)
