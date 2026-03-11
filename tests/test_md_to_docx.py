from pathlib import Path
import pytest
from docx import Document
from md_docx_converter.md_to_docx import convert_md_to_docx

FIXTURES = Path(__file__).parent / "fixtures"
TEMPLATE = Path(r"C:\Users\Chris\AppData\Roaming\Microsoft\Templates\Normal.dotm")

if not TEMPLATE.exists():
    pytest.skip("Word Normal template not found — skipping DOCX tests", allow_module_level=True)


def _convert(md_file):
    out = FIXTURES / (md_file.stem + "_out.docx")
    convert_md_to_docx(md_file, out, TEMPLATE)
    return Document(str(out))


def test_single_h1_becomes_title():
    doc = _convert(FIXTURES / "simple.md")
    styles = [p.style.name for p in doc.paragraphs if p.text.strip()]
    assert styles[0] == "Title"


def test_h2_becomes_heading1_when_title_present():
    doc = _convert(FIXTURES / "simple.md")
    styles = [p.style.name for p in doc.paragraphs if p.text.strip()]
    assert "Heading 1" in styles


def test_bold_run_is_bold():
    doc = _convert(FIXTURES / "simple.md")
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text == "bold":
                assert run.bold
                return
    raise AssertionError("No bold run found")


def test_italic_run_is_italic():
    doc = _convert(FIXTURES / "simple.md")
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text == "italic":
                assert run.italic
                return
    raise AssertionError("No italic run found")


def test_blockquote_uses_quote_style():
    doc = _convert(FIXTURES / "simple.md")
    quote_paras = [p for p in doc.paragraphs if p.style.name == "Quote"]
    assert len(quote_paras) >= 1
    assert "blockquote" in quote_paras[0].text


def test_table_created():
    doc = _convert(FIXTURES / "simple.md")
    assert len(doc.tables) == 1


def test_table_has_correct_content():
    doc = _convert(FIXTURES / "simple.md")
    table = doc.tables[0]
    assert table.cell(0, 0).text == "Col A"
    assert table.cell(1, 0).text == "One"


def test_image_embedded():
    """Image referenced in MD is embedded in DOCX."""
    md_content = "# Title\n\n![alt](sample.png)\n"
    import tempfile, shutil
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        shutil.copy(FIXTURES / "sample.png", tmp / "sample.png")
        md_path = tmp / "test_img.md"
        md_path.write_text(md_content, encoding="utf-8")
        out_path = tmp / "test_img.docx"
        convert_md_to_docx(md_path, out_path, TEMPLATE)
        doc = Document(str(out_path))
        body_xml = doc.element.body.xml
        assert "graphicData" in body_xml or "drawing" in body_xml
