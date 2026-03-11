"""Run once to create test fixture DOCX files. Re-run after adding new sections."""
from docx import Document
from docx.shared import Inches
import os

fixtures = os.path.dirname(os.path.abspath(__file__))

# with_title.docx — has Title + Heading 1
doc = Document()
doc.add_paragraph("My Document Title", style="Title")
doc.add_paragraph("Section One", style="Heading 1")
doc.add_paragraph("Body text here.")
doc.save(os.path.join(fixtures, "with_title.docx"))

# no_title.docx — only Heading 1, no Title
doc2 = Document()
doc2.add_paragraph("Section One", style="Heading 1")
doc2.add_paragraph("Section Two", style="Heading 1")
doc2.add_paragraph("Body text here.")
doc2.save(os.path.join(fixtures, "no_title.docx"))

# with_image.docx — has an embedded image (requires sample.png to exist first)
sample_png = os.path.join(fixtures, "sample.png")
if os.path.exists(sample_png):
    doc3 = Document()
    doc3.add_paragraph("Before image")
    doc3.add_picture(sample_png, width=Inches(1))
    doc3.add_paragraph("After image")
    doc3.save(os.path.join(fixtures, "with_image.docx"))
    print("with_image.docx created")
else:
    print("WARNING: sample.png not found — skipping with_image.docx")

# rich.docx — tests all element types for DOCX→MD
doc4 = Document()
doc4.add_paragraph("Document Title", style="Title")
doc4.add_paragraph("Section One", style="Heading 1")
doc4.add_paragraph("Normal body text.")
p = doc4.add_paragraph()
run = p.add_run("bold word")
run.bold = True
p.add_run(" and plain.")
doc4.add_paragraph("Quote text", style="Quote")
table = doc4.add_table(rows=2, cols=2)
table.cell(0, 0).text = "H1"
table.cell(0, 1).text = "H2"
table.cell(1, 0).text = "R1"
table.cell(1, 1).text = "R2"
doc4.save(os.path.join(fixtures, "rich.docx"))

print("All fixtures created.")
